#!/usr/bin/env python3
"""
Round‑trip Word ↔ Markdown with style preservation.

Usage:
  # DOCX -> MD
  python docx_md_roundtrip.py to-md "input.docx" -o out.md --media-dir media

  # MD -> DOCX (use the original DOCX as reference to keep the exact styles)
  python docx_md_roundtrip.py to-docx out.md -o new.docx --ref "input.docx"

Notes:
- Requires: python-docx, mammoth, pyyaml, and pandoc (CLI) on PATH.
- Preserves paragraph, character, and table styles via Markdown attributes like:
      {custom-style="Heading 2"}
- Images export to --media-dir on DOCX -> MD and are re-linked in the MD.
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile
import uuid
from pathlib import Path
from typing import Dict, Tuple

import yaml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
import mammoth


# ---------- helpers ----------

def check_pandoc() -> None:
    try:
        subprocess.run(
            ["pandoc", "--version"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=True,
        )
    except Exception:
        sys.exit("pandoc not found on PATH. Please install pandoc and try again.")


def pandoc_heading_arg() -> str:
    """Return the appropriate pandoc flag for ATX-style headings."""
    try:
        p = subprocess.run(
            ["pandoc", "--version"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=True,
            text=True,
        )
        version = p.stdout.splitlines()[0].split()[1]
        major = int(version.split(".")[0])
        if major >= 3:
            return "--markdown-headings=atx"
    except Exception:
        pass
    return "--atx-headers"


def slug_token(name: str) -> str:
    """
    Convert a Word style name into a safe token we can round‑trip through HTML classes.
    """
    token = re.sub(r"[^A-Za-z0-9]+", "_", name.strip())
    token = re.sub(r"_+", "_", token)
    return token.strip("_") or "Style"


def collect_used_styles(docx_path: Path) -> Tuple[Dict[str, str], Dict[str, str], Dict[str, str]]:
    """Return three dicts: {token: original_name} for paragraph, character, table styles used in the document."""
    doc = Document(str(docx_path))
    p_styles: Dict[str, str] = {}
    r_styles: Dict[str, str] = {}
    t_styles: Dict[str, str] = {}

    # paragraphs + runs
    for p in doc.paragraphs:
        if p.style and p.style.type == WD_STYLE_TYPE.PARAGRAPH:
            n = p.style.name
            p_styles.setdefault(slug_token(n), n)
        for r in p.runs:
            try:
                if r.style and r.style.type == WD_STYLE_TYPE.CHARACTER:
                    n = r.style.name
                    r_styles.setdefault(slug_token(n), n)
            except Exception:
                # Some runs may not expose style cleanly; ignore.
                pass

    # tables
    for t in doc.tables:
        try:
            if t.style and t.style.type == WD_STYLE_TYPE.TABLE:
                n = t.style.name
                t_styles.setdefault(slug_token(n), n)
        except Exception:
            pass

    return p_styles, r_styles, t_styles


def build_mammoth_style_map(p_styles, r_styles, t_styles) -> str:
    """
    Build a Mammoth style map string that:
      - Maps Word heading paragraph styles to h1..h6 and attaches token classes
      - Maps all other paragraph styles to p.token
      - Maps character styles to span.token
      - Maps table styles to table.token
    """
    lines = []

    # Headings
    for level in range(1, 7):
        w_name = f"Heading {level}"
        token = slug_token(w_name)
        # Include headings even if not detected (defensive)
        lines.append(f'p[style-name="{w_name}"] => h{level}.{token}')

    # Other paragraph styles
    for token, name in p_styles.items():
        if re.fullmatch(r"Heading [1-6]", name):
            continue  # already handled
        lines.append(f'p[style-name="{name}"] => p.{token}')

    # Character styles
    for token, name in r_styles.items():
        lines.append(f'r[style-name="{name}"] => span.{token}')

    # Table styles
    for token, name in t_styles.items():
        lines.append(f'table[style-name="{name}"] => table.{token}')

    return "\n".join(lines)


def write_lua_filter(dest: Path) -> Path:
    """
    Lua filter that:
      - Reads metadata.style_map (token -> Word style name)
      - For any Header/Para/Div/Span/Table with a class that matches a token,
        sets attr.attributes["custom-style"] = original style name
        and removes the token from classes (cosmetic).
      - If custom-style already exists, leaves it alone.
    """
    lua = r'''
local style_map = {}

function Meta(meta)
  if meta["style_map"] then
    for k, v in pairs(meta["style_map"]) do
      style_map[k] = pandoc.utils.stringify(v)
    end
  end
  return nil
end

local function apply_custom_style(attr)
  if attr and attr.attributes and attr.attributes["custom-style"] then
    return attr
  end
  if attr and attr.classes then
    local to_remove = {}
    for i, cls in ipairs(attr.classes) do
      local mapped = style_map[cls]
      if mapped and mapped ~= "" then
        attr.attributes = attr.attributes or {}
        attr.attributes["custom-style"] = mapped
        table.insert(to_remove, i)
        break
      end
    end
    table.sort(to_remove, function(a,b) return a>b end)
    for _, idx in ipairs(to_remove) do
      table.remove(attr.classes, idx)
    end
  end
  return attr
end

function Header(el) el.attr = apply_custom_style(el.attr); return el end
function Para(el)   el.attr = apply_custom_style(el.attr); return el end
function Div(el)    el.attr = apply_custom_style(el.attr); return el end
function Span(el)   el.attr = apply_custom_style(el.attr); return el end
function Table(el)  el.attr = apply_custom_style(el.attr); return el end
'''
    dest.write_text(lua, encoding="utf-8")
    return dest


def docx_to_md(input_docx: Path, out_md: Path, media_dir: Path) -> Path:
    check_pandoc()
    media_dir.mkdir(parents=True, exist_ok=True)

    # 1) Collect used styles
    p_styles, r_styles, t_styles = collect_used_styles(input_docx)

    # 2) Build Mammoth style map (so HTML gets classes for styles)
    style_map_text = build_mammoth_style_map(p_styles, r_styles, t_styles)

    # 3) Convert DOCX -> HTML with mammoth; extract images
    def save_image(image):
        # e.g., "image/png"; handle missing content-type safely
        ct = (image.content_type or "image/png").lower()
        ext_map = {
            "image/jpeg": "jpg",
            "image/jpg": "jpg",
            "image/png": "png",
            "image/gif": "gif",
            "image/tiff": "tif",
            "image/bmp": "bmp",
            "image/svg+xml": "svg",
            "image/x-emf": "emf",
            "image/x-wmf": "wmf",
        }
        ext = ext_map.get(ct, ct.split("/")[-1] or "png")

        fname = f"img-{uuid.uuid4().hex}.{ext}"
        target = media_dir / fname
        target.parent.mkdir(parents=True, exist_ok=True)

        # IMPORTANT: read inside the context manager (fixes 'closing' object error)
        with image.open() as img_bytes, target.open("wb") as out:
            out.write(img_bytes.read())

        # return the relative src used in HTML
        return {"src": str(Path(os.path.relpath(target, out_md.parent)).as_posix())}

    with open(input_docx, "rb") as f:
        html_result = mammoth.convert_to_html(
            f,
            style_map=style_map_text,
            convert_image=mammoth.images.inline(save_image),
        )
    html = html_result.value

    # 4) Prepare metadata (style tokens -> real names) for the Lua filter
    metadata = {"style_map": {**p_styles, **r_styles, **t_styles}}
    tmpdir = Path(tempfile.mkdtemp(prefix="roundtrip-"))
    meta_file = tmpdir / "meta.yaml"
    meta_file.write_text(
        yaml.safe_dump(metadata, sort_keys=True, allow_unicode=True),
        encoding="utf-8",
    )
    lua_filter = write_lua_filter(tmpdir / "classes_to_customstyle.lua")

    # 5) HTML -> Markdown (keep attributes); Lua filter injects custom-style
    out_md.parent.mkdir(parents=True, exist_ok=True)
    cmd = [
        "pandoc",
        "--from=html",
        "--to=markdown+bracketed_spans+fenced_divs+pipe_tables+header_attributes",
        "--wrap=none",
        pandoc_heading_arg(),
        f"--metadata-file={meta_file}",
        f"--lua-filter={lua_filter}",
        "-o",
        str(out_md),
    ]
    p = subprocess.run(cmd, input=html, text=True, capture_output=True)
    if p.returncode != 0:
        raise RuntimeError(f"pandoc HTML->MD failed:\n{p.stderr}")

    # 6) Prepend YAML header (keeps the mapping in the MD file for future edits)
    md_text = out_md.read_text(encoding="utf-8")
    front_matter = "---\n" + yaml.safe_dump(metadata, sort_keys=True, allow_unicode=True) + "---\n\n"
    out_md.write_text(front_matter + md_text, encoding="utf-8")

    shutil.rmtree(tmpdir, ignore_errors=True)
    return out_md


def md_to_docx(in_md: Path, out_docx: Path, reference_docx: Path | None = None) -> Path:
    check_pandoc()

    tmpdir = Path(tempfile.mkdtemp(prefix="roundtrip-"))
    lua_filter = write_lua_filter(tmpdir / "classes_to_customstyle.lua")

    cmd = [
        "pandoc",
        str(in_md),
        "--from=markdown+bracketed_spans+fenced_divs+pipe_tables+header_attributes",
        "--to=docx",
        "--wrap=none",
        f"--lua-filter={lua_filter}",
        "-o",
        str(out_docx),
    ]
    if reference_docx:
        # ensure the reference file exists (nicer error than pandoc's)
        if not Path(reference_docx).exists():
            shutil.rmtree(tmpdir, ignore_errors=True)
            raise FileNotFoundError(f"Reference DOCX not found: {reference_docx}")
        cmd.insert(-2, f"--reference-doc={reference_docx}")  # before -o

    p = subprocess.run(cmd, capture_output=True, text=True)
    if p.returncode != 0:
        shutil.rmtree(tmpdir, ignore_errors=True)
        raise RuntimeError(f"pandoc MD->DOCX failed:\n{p.stderr}")

    shutil.rmtree(tmpdir, ignore_errors=True)
    return out_docx


# ---------- CLI ----------

def main() -> None:
    ap = argparse.ArgumentParser(description="Round-trip DOCX <-> MD with style preservation.")
    sub = ap.add_subparsers(dest="cmd", required=True)

    to_md = sub.add_parser("to-md", help="Convert DOCX to Markdown (preserving Word styles).")
    to_md.add_argument("input", type=Path)
    to_md.add_argument("-o", "--out", type=Path, required=True, help="Output .md")
    to_md.add_argument("--media-dir", type=Path, default=Path("media"), help="Relative path for exported images")

    to_docx = sub.add_parser("to-docx", help="Convert Markdown back to DOCX (re-applying Word styles).")
    to_docx.add_argument("input", type=Path, help="Input .md")
    to_docx.add_argument("-o", "--out", type=Path, required=True, help="Output .docx")
    to_docx.add_argument("--ref", type=Path, default=None, help="Reference .docx with style definitions (recommended: the original DOCX)")

    args = ap.parse_args()

    if args.cmd == "to-md":
        docx_to_md(args.input, args.out, args.media_dir)
        print(f"Wrote Markdown: {args.out}")
    else:
        md_to_docx(args.input, args.out, args.ref)
        print(f"Wrote DOCX: {args.out}")


if __name__ == "__main__":
    main()
